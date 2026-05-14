from io import BytesIO

from docx import Document

from app.rag.extractors.base_extractor import BaseExtractor
from app.rag.models.raw_document import RawDocument


class DocxExtractor(BaseExtractor):
    """
    DOCX extractor.

    Supports:
    - byte-stream ingestion
    - local filesystem fallback

    Future:
    - object-storage streaming
    """

    def extract_text(
        self,
        raw_document: RawDocument,
    ) -> str:

        if not raw_document:
            raise ValueError("raw_document is required")

        # Preferred distributed-safe extraction
        if raw_document.content:

            document = Document(
                BytesIO(raw_document.content)
            )

        # Local dev fallback
        elif raw_document.local_path:

            document = Document(
                raw_document.local_path
            )

        # Future: storage_uri stream extraction
        elif raw_document.storage_uri:

            raise NotImplementedError(
                "storage_uri extraction not implemented yet"
            )

        else:

            raise ValueError(
                f"No content available for file: "
                f"{raw_document.file_name}"
            )

        extracted_parts = []

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:
                extracted_parts.append(text)

        for table in document.tables:

            for row in table.rows:

                row_values = []

                for cell in row.cells:

                    cell_text = cell.text.strip()

                    if cell_text:
                        row_values.append(cell_text)

                if row_values:

                    extracted_parts.append(
                        " | ".join(row_values)
                    )

        extracted_text = (
            "\n".join(extracted_parts).strip()
        )

        if not extracted_text:

            return (
                f"[WARN] No readable text extracted "
                f"from DOCX file: "
                f"{raw_document.file_name}"
            )

        return extracted_text